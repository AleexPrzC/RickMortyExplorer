from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "RickMortyExplorer_Codelab_Academico.docx"


ACCENT = "167A72"
ACCENT_DARK = "0D4F4A"
MUTED = "667085"
LIGHT = "E6F4F1"
TABLE_HEADER = "D7F0EB"
CODE_BG = "F4F6F8"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D0D5DD"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_pct=95):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_pct * 50))
    tbl_w.set(qn("w:type"), "pct")


def add_page_number(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("RickMorty Explorer · Pagina ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(102, 112, 133)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    add_page_number(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Title", 28, ACCENT_DARK),
        ("Heading 1", 18, ACCENT_DARK),
        ("Heading 2", 13, ACCENT),
        ("Heading 3", 11.5, MUTED),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    r = p.add_run("RickMorty Explorer")
    r.font.name = "Aptos Display"
    r.font.size = Pt(32)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(ACCENT_DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Codelab completo y memoria tecnica de apoyo")
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor.from_string(ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("Proyecto Android en Java + XML con Retrofit, Gson, RxJava2 y RecyclerView")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 76)
    data = [
        ("Asignatura", "Computacion Movil"),
        ("Proyecto", "RickMortyExplorer"),
        ("API REST", "Rick and Morty API"),
        ("Tecnologias", "Java, XML, Retrofit, Gson, RxJava2, RecyclerView, Picasso"),
        ("Fecha", "Mayo de 2026"),
    ]
    for row, (label, value) in zip(table.rows, data):
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], TABLE_HEADER)
        row.cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_page_break()


def add_toc(doc):
    doc.add_heading("Indice", level=1)
    items = [
        "1. Objetivo del proyecto",
        "2. Requisitos cubiertos",
        "3. API utilizada",
        "4. Estructura del proyecto",
        "5. Configuracion Gradle",
        "6. AndroidManifest",
        "7. Recursos XML y layouts",
        "8. Modelos Java",
        "9. Retrofit y endpoints",
        "10. RecyclerView: Adapter y ViewHolder",
        "11. MainActivity",
        "12. DetailActivity",
        "13. Flujo completo de la aplicacion",
        "14. Estados de loading y error",
        "15. Errores tipicos y solucion",
        "16. Pruebas realizadas",
        "17. Posibles ampliaciones",
        "18. Guia de defensa",
    ]
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
    doc.add_page_break()


def add_note(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 95)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    set_cell_border(cell, "B6E2D9")
    set_cell_margins(cell, 140, 160, 140, 160)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
    p.add_run("\n" + text)


def add_code(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 95)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CODE_BG)
    set_cell_border(cell, "E4E7EC")
    set_cell_margins(cell, 120, 140, 120, 140)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(code.strip())
    r.font.name = "Consolas"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    r.font.size = Pt(8.5)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def read_file(relative):
    return (ROOT / relative).read_text(encoding="utf-8")


def main():
    doc = Document()
    style_document(doc)
    add_cover(doc)
    add_toc(doc)

    doc.add_heading("1. Objetivo del proyecto", level=1)
    doc.add_paragraph(
        "El objetivo de este proyecto es desarrollar una aplicacion Android completa, "
        "sencilla y defendible que consuma una API REST externa y muestre informacion "
        "en una interfaz construida con Java y XML."
    )
    add_note(
        doc,
        "Idea principal",
        "RickMorty Explorer muestra una lista de personajes de la Rick and Morty API. "
        "Al seleccionar un personaje, se abre una pantalla de detalle con datos ampliados."
    )

    doc.add_heading("2. Requisitos cubiertos", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 95)
    headers = ["Requisito", "Implementacion", "Estado"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, TABLE_HEADER)
        set_cell_border(cell)
        set_cell_margins(cell)
        cell.paragraphs[0].runs[0].font.bold = True
    rows = [
        ("Java + Views/XML", "Activities en Java y layouts XML", "Cumplido"),
        ("Retrofit + Gson + RxJava2", "RetrofitClient, GsonConverterFactory y RxJava2CallAdapterFactory", "Cumplido"),
        ("RecyclerView", "CharacterAdapter con CharacterViewHolder", "Cumplido"),
        ("Endpoint lista", "GET /api/character?page=1", "Cumplido"),
        ("Endpoint detalle", "GET /api/character/{id}", "Cumplido"),
        ("Loading y error", "ProgressBar, contenedor de error y boton de reintento", "Cumplido"),
    ]
    for values in rows:
        row = table.add_row()
        for i, value in enumerate(values):
            cell = row.cells[i]
            cell.text = value
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_heading("3. API utilizada", level=1)
    doc.add_paragraph("Se utiliza la Rick and Morty API, una API publica sin clave de acceso.")
    add_code(
        doc,
        """
Base URL: https://rickandmortyapi.com/api/
Lista:    GET character?page=1
Detalle:  GET character/{id}
        """,
    )
    doc.add_paragraph(
        "La respuesta de lista contiene un objeto con informacion de paginacion y un array "
        "`results`. Cada elemento de `results` representa un personaje. La respuesta de detalle "
        "devuelve un unico personaje con campos como id, name, status, species, gender, origin, "
        "location e image."
    )

    doc.add_heading("4. Estructura del proyecto", level=1)
    add_code(
        doc,
        """
com.example.rickmortyexplorer
|-- model
|   |-- CharacterItem.java
|   |-- CharacterListResponse.java
|   |-- Info.java
|   `-- Origin.java
|-- rest
|   |-- RetrofitClient.java
|   `-- api
|       `-- RickMortyApiService.java
|-- ui
|   |-- MainActivity.java
|   |-- DetailActivity.java
|   `-- adapter
|       `-- CharacterAdapter.java
`-- util
    `-- CharacterUtils.java
        """,
    )
    doc.add_paragraph(
        "Esta separacion ayuda a que cada parte tenga una responsabilidad clara: modelos para datos, "
        "rest para red, ui para pantallas, adapter para RecyclerView y util para formateo."
    )

    doc.add_heading("5. Configuracion Gradle", level=1)
    doc.add_paragraph("El proyecto usa Gradle con Version Catalogs para centralizar versiones.")
    add_code(doc, read_file("gradle/libs.versions.toml"))
    doc.add_paragraph("En el modulo app se activa ViewBinding y se declaran las dependencias principales.")
    add_code(doc, read_file("app/build.gradle.kts"))

    doc.add_heading("6. AndroidManifest", level=1)
    doc.add_paragraph(
        "El Manifest declara el permiso de Internet, necesario para consumir la API REST, "
        "y registra las dos actividades de la aplicacion."
    )
    add_code(doc, read_file("app/src/main/AndroidManifest.xml"))

    doc.add_heading("7. Recursos XML y layouts", level=1)
    doc.add_heading("7.1 Pantalla principal", level=2)
    doc.add_paragraph(
        "El layout `activity_main.xml` contiene una cabecera, un RecyclerView, un ProgressBar "
        "y un contenedor de error con boton de reintento."
    )
    add_code(doc, read_file("app/src/main/res/layout/activity_main.xml")[:3500] + "\n...")
    doc.add_heading("7.2 Fila del RecyclerView", level=2)
    doc.add_paragraph(
        "La fila `row_character.xml` muestra imagen, nombre, especie, genero y estado del personaje."
    )
    add_code(doc, read_file("app/src/main/res/layout/row_character.xml"))
    doc.add_heading("7.3 Pantalla de detalle", level=2)
    doc.add_paragraph(
        "El layout `activity_detail.xml` repite el mismo enfoque de estados: contenido, carga y error."
    )

    doc.add_heading("8. Modelos Java", level=1)
    doc.add_paragraph(
        "Los modelos son POJOs que coinciden con el JSON devuelto por la API. Gson rellena estos "
        "objetos automaticamente."
    )
    add_code(doc, read_file("app/src/main/java/com/example/rickmortyexplorer/model/CharacterItem.java"))
    add_code(doc, read_file("app/src/main/java/com/example/rickmortyexplorer/model/CharacterListResponse.java"))

    doc.add_heading("9. Retrofit y endpoints", level=1)
    doc.add_paragraph(
        "`RetrofitClient` centraliza la URL base y registra Gson y RxJava2. "
        "`RickMortyApiService` define los endpoints mediante anotaciones."
    )
    add_code(doc, read_file("app/src/main/java/com/example/rickmortyexplorer/rest/RetrofitClient.java"))
    add_code(doc, read_file("app/src/main/java/com/example/rickmortyexplorer/rest/api/RickMortyApiService.java"))

    doc.add_heading("10. RecyclerView: Adapter y ViewHolder", level=1)
    doc.add_paragraph(
        "El adaptador recibe una lista de personajes y se encarga de pintar cada fila. "
        "Tambien define una interfaz de click para abrir el detalle."
    )
    add_code(doc, read_file("app/src/main/java/com/example/rickmortyexplorer/ui/adapter/CharacterAdapter.java"))

    doc.add_heading("11. MainActivity", level=1)
    doc.add_paragraph(
        "La pantalla principal crea el servicio Retrofit, configura el RecyclerView y llama al "
        "endpoint de lista. La peticion se ejecuta en segundo plano con `Schedulers.io()` y vuelve "
        "al hilo principal con `AndroidSchedulers.mainThread()`."
    )
    add_code(doc, read_file("app/src/main/java/com/example/rickmortyexplorer/ui/MainActivity.java"))

    doc.add_heading("12. DetailActivity", level=1)
    doc.add_paragraph(
        "La pantalla de detalle recibe el id del personaje por Intent y consume el endpoint de detalle."
    )
    add_code(doc, read_file("app/src/main/java/com/example/rickmortyexplorer/ui/DetailActivity.java"))

    doc.add_heading("13. Flujo completo de la aplicacion", level=1)
    add_bullets(
        doc,
        [
            "El usuario abre la app.",
            "MainActivity muestra un ProgressBar.",
            "Se llama a GET character?page=1.",
            "Cuando llegan los datos, el RecyclerView muestra los personajes.",
            "El usuario pulsa un personaje.",
            "DetailActivity recibe el id por Intent.",
            "Se llama a GET character/{id}.",
            "La pantalla de detalle muestra la informacion ampliada.",
        ],
    )

    doc.add_heading("14. Estados de loading y error", level=1)
    doc.add_paragraph(
        "La app no deja la pantalla vacia durante las peticiones. Tanto la lista como el detalle "
        "tienen funciones para mostrar carga, contenido o error."
    )
    add_code(
        doc,
        """
showLoading()  -> muestra ProgressBar y oculta contenido/error
showContent()  -> muestra la informacion cargada
showError()    -> muestra mensaje de error y boton de reintento
        """,
    )

    doc.add_heading("15. Errores tipicos y solucion", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 95)
    for i, h in enumerate(["Problema", "Solucion"]):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, TABLE_HEADER)
        set_cell_border(cell)
        set_cell_margins(cell)
        cell.paragraphs[0].runs[0].font.bold = True
    for problem, solution in [
        ("La app no tiene Internet", "Comprobar el permiso INTERNET en AndroidManifest.xml."),
        ("No aparecen los bindings", "Verificar que viewBinding = true y que los ids existen en XML."),
        ("Error de Gradle", "Sincronizar proyecto y revisar versiones de Gradle/AGP."),
        ("No carga la API", "Comprobar conexion del emulador y URL base de Retrofit."),
        ("No abre detalle", "Comprobar que se envia EXTRA_CHARACTER_ID en el Intent."),
    ]:
        row = table.add_row()
        for i, text in enumerate([problem, solution]):
            cell = row.cells[i]
            cell.text = text
            set_cell_border(cell)
            set_cell_margins(cell)

    doc.add_heading("16. Pruebas realizadas", level=1)
    add_bullets(
        doc,
        [
            "Sincronizacion Gradle correcta.",
            "Compilacion debug correcta.",
            "Ejecucion en emulador Android.",
            "Carga de lista de personajes.",
            "Scroll en RecyclerView.",
            "Apertura de pantalla de detalle al pulsar un personaje.",
        ],
    )

    doc.add_heading("17. Posibles ampliaciones", level=1)
    add_bullets(
        doc,
        [
            "Anadir buscador por nombre.",
            "Anadir paginacion para cargar mas personajes.",
            "Guardar favoritos localmente.",
            "Migrar a MVVM.",
            "Usar fragments para navegacion adaptable.",
            "Mostrar localizaciones o episodios como pantallas adicionales.",
        ],
    )

    doc.add_heading("18. Guia de defensa", level=1)
    doc.add_paragraph("Para defender el proyecto se recomienda explicar el flujo con este orden:")
    add_bullets(
        doc,
        [
            "Primero, explicar la API y los dos endpoints usados.",
            "Despues, ensenar RetrofitClient y RickMortyApiService.",
            "A continuacion, mostrar MainActivity y el RecyclerView.",
            "Luego, abrir CharacterAdapter para identificar Adapter y ViewHolder.",
            "Por ultimo, mostrar DetailActivity y explicar el segundo endpoint.",
        ],
    )
    add_note(
        doc,
        "Resumen final",
        "RickMorty Explorer cumple los requisitos minimos del proyecto final y queda preparado "
        "para ampliaciones opcionales sin cambiar la base arquitectonica."
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
